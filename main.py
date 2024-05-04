########################### Start of imports ###########################

import requests
import msal
import openai
import asyncio
import win32com.client # pip install pywin32
import datetime
import os
import glob
import re

from docx import Document
from docx.shared import Pt
from oauthlib.oauth2 import WebApplicationClient
from msgraph import GraphServiceClient
#from azure.identity import ClientSecret
from azure.identity import ClientSecretCredential
from pathlib import Path
from datetime import date
from PyQt6 import QtCore, QtGui, QtWidgets

########################### End of imports ###########################
#
#
#
#
########################### Start of helper functions ###########################

global emailGlobal 
global dateGlobal 
global subjectGlobal 
global dateFromG
global subjectG
global emailWG

emailGlobal = False
dateGlobal = False
subjectGlobal = False
dateFromG = date.today()
subjectG = "default"
emailWG = "default"

##################################UI CODE##################################
class Ui_Dialog(object):
    def setupUi(self, Dialog):
        Dialog.setObjectName("Dialog")
        Dialog.resize(1117, 808)
        Dialog.setStyleSheet("background-color: rgb(179, 132, 51)")
        self.label = QtWidgets.QLabel(parent=Dialog)
        self.label.setGeometry(QtCore.QRect(270, 30, 631, 121))
        self.label.setStyleSheet("font-family: Arial;\n"
                                  "font-size: 80px;\n"
                                  "font-weight: bold;\n"
                                  " qproperty-alignment: \'AlignCenter\';")
        self.label.setObjectName("label")

        self.checkBox = QtWidgets.QCheckBox(parent=Dialog)
        self.checkBox.setGeometry(QtCore.QRect(120, 310, 881, 71))
        self.checkBox.setStyleSheet("font-family: Arial;\n"
                                     "font-size: 30px;\n"
                                     "font-weight: bold;\n"
                                     )
        self.checkBox.setObjectName("checkBox")

        self.checkBox_2 = QtWidgets.QCheckBox(parent=Dialog)
        self.checkBox_2.setGeometry(QtCore.QRect(120, 210, 941, 51))
        self.checkBox_2.setStyleSheet("font-family: Arial;\n"
                                       "font-size: 30px;\n"
                                       "font-weight: bold;\n"
                                      )
        self.checkBox_2.setObjectName("checkBox_2")

        self.checkBox_3 = QtWidgets.QCheckBox(parent=Dialog)
        self.checkBox_3.setGeometry(QtCore.QRect(120, 265, 941, 51))
        self.checkBox_3.setStyleSheet("font-family: Arial;\n"
                                       "font-size: 30px;\n"
                                       "font-weight: bold;\n"
                                       )
        self.checkBox_3.setObjectName("checkBox_3")

        self.lineEdit = QtWidgets.QLineEdit(parent=Dialog)
        self.lineEdit.setGeometry(QtCore.QRect(90, 490, 451, 25))
        self.lineEdit.setStyleSheet("background-color:rgb(0, 0, 0)\n"
                                    "\n"
                                    "\n"
                                    "")
        self.lineEdit.setObjectName("lineEdit")
        self.dateEdit = QtWidgets.QDateEdit(parent=Dialog)
        self.dateEdit.setGeometry(QtCore.QRect(630, 480, 370, 51))
        self.dateEdit.setStyleSheet("font-family: Arial;\n"
                                     "font-size: 40px;\n"
                                     "font-weight: bold;\n"
                                     " qproperty-alignment: \'AlignCenter\';")
        self.dateEdit.setObjectName("dateEdit")
        self.label_2 = QtWidgets.QLabel(parent=Dialog)
        self.label_2.setGeometry(QtCore.QRect(90, 450, 451, 41))
        self.label_2.setStyleSheet("font-family: Arial;\n"
                                    "font-size: 20px;\n"
                                    "font-weight: bold;\n"
                                    " qproperty-alignment: \'AlignCenter\';")
        self.label_2.setObjectName("label_2")
        self.pushButton = QtWidgets.QPushButton(parent=Dialog)
        self.pushButton.setGeometry(QtCore.QRect(350, 660, 431, 91))
        self.pushButton.setStyleSheet("font-family: Arial;\n"
                                       "font-size: 40px;\n"
                                       "font-weight: bold;\n"
                                       )
        self.pushButton.setObjectName("pushButton")
        self.label_3 = QtWidgets.QLabel(parent=Dialog)
        self.label_3.setGeometry(QtCore.QRect(650, 450, 311, 21))
        self.label_3.setStyleSheet("font-family: Arial;\n"
                                    "font-size: 20px;\n"
                                    "font-weight: bold;\n"
                                    " qproperty-alignment: \'AlignCenter\';")
        self.label_3.setObjectName("label_3")
        self.lineEdit_2 = QtWidgets.QLineEdit(parent=Dialog)
        self.lineEdit_2.setGeometry(QtCore.QRect(90, 570, 451, 25))
        self.lineEdit_2.setStyleSheet("background-color:rgb(0, 0, 0)\n"
                                    "\n"
                                    "\n"
                                    "")
        self.lineEdit_2.setObjectName("lineEdit_2")

        self.label_4 = QtWidgets.QLabel(parent=Dialog)
        self.label_4.setGeometry(QtCore.QRect(90, 530, 451, 41))
        self.label_4.setStyleSheet("font-family: Arial;\n"
                                    "font-size: 20px;\n"
                                    "font-weight: bold;\n"
                                    " qproperty-alignment: \'AlignCenter\';")
        self.label_4.setObjectName("label_4")

        self.retranslateUi(Dialog)
        QtCore.QMetaObject.connectSlotsByName(Dialog)

    def retranslateUi(self, Dialog):
        _translate = QtCore.QCoreApplication.translate
        Dialog.setWindowTitle(_translate("Dialog", "Dialog"))
        Dialog.setWhatsThis(_translate("Dialog", "<html><head/><body><p>Dispute Lens</p></body></html>"))
        self.label.setText(_translate("Dialog", "DisputeLens"))
        self.checkBox.setText(_translate("Dialog", "Generate Timeline from specific date [2]"))
        self.checkBox_2.setText(_translate("Dialog", "Generate Timeline with specific email address [1]"))
        self.checkBox_3.setText(_translate("Dialog", "Generate Timeline with specific email subject line [3]"))
        self.label_2.setText(_translate("Dialog", "Enter receivers email address [1]"))
        self.pushButton.setText(_translate("Dialog", "Generate Timeline"))
        self.label_3.setText(_translate("Dialog", "Enter the specific date [2]"))
        self.label_4.setText(_translate("Dialog", "Enter email subject line [3]"))

##################################END OF UI CODE##################################

def txt_to_docx(txt_file, docx_file):
    # Create a new Document
    doc = Document()
    # Open the .txt file
    with open(txt_file, 'r', encoding='utf-8') as file:
        # Read the file
        text = file.read()
        # Add the text to the Document
        doc.add_paragraph(text)
        # Save the Document
        doc.save(docx_file)

    return

# read the email thread from emails_raw.docx and return the content as a string
def read_document(docx_file_name):
    # Open the existing document
    doc = Document(docx_file_name)

    # Read the content of the document
    content = []
    for para in doc.paragraphs:
        content.append(para.text)

    # Join all the paragraphs together
    string_content = ' '.join(content)

    return string_content

# Create a new document with any string argument
def create_document(string_content, docx_name):
    # Create a new document
    doc = Document()
    # Add content to the document
    doc.add_paragraph(string_content)
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)  # Use Pt for font size

    # Save the document
    doc.save(docx_name)

    return

def write_to_existing_document(string_content, docx_name):
    # Open the existing document
    doc = Document(docx_name)
    
    # Add content to the document
    doc.add_paragraph(string_content)
    
    # Save the document
    doc.save(docx_name)

    return

def getGPT3Response(prompt):
    # Set your OpenAI API key
    openai.api_key = 'sk-proj-UGMA2Nv0TPrVSCpl0dngT3BlbkFJYgd9SSZ6dV2avNJLQt0C'
    # Make a GPT-3.5 API call
    response = openai.ChatCompletion.create(
    model="gpt-3.5-turbo",
    messages=[
            {"role": "user", "content": prompt},
        ]
    )
    # Get the content from the GPT-3 response
    string_content = response['choices'][0]['message']['content']

    return string_content

def getEmails(givenSubject, emailWith):
    # Convert the given date to a datetime object
    # if isinstance(givenDateFrom, str):
    #     givenDateFrom = datetime.strptime(givenDateFrom, '%Y-%m-%d').date()

    # Create output folder
    output_dir = Path.cwd() / 'output'
    output_dir.mkdir(parents=True,exist_ok=True)

    # Empty the output folder
    files = glob.glob(str(output_dir / '*'))
    for f in files:
        os.remove(f)

    #Connect to Outlook
    outlook = win32com.client.Dispatch("Outlook.Application").GetNamespace("MAPI")

    #Connect to folder
    inbox = outlook.GetDefaultFolder(6) #6 is the inbox folder ID

    #Get all emails in the inboxqq
    messages = inbox.Items

    #Sort the emails by ReceivedTime, in descending order
    messages.Sort("[ReceivedTime]", True)

    return messages

def formTimeline(messages, dateString):
    num_messages = len(messages)
    print(f"Number of messages: {num_messages}")

    for message in messages:
        print(num_messages)
        num_messages -= 1
        # Check if the item is an email message
        if hasattr(message, "MessageClass") and message.MessageClass.startswith("IPM.Note"):
            
            # Get recipients email address
            to_email_address = "default"
            body2 = message.Body
            # Get the line that starts with 'To:'
            to_line = next((line for line in body2.split('\n') if line.startswith('To:')), None)

            if to_line is not None:
                # Use a regular expression to find the email address
                match = re.search(r'<(.+)>', to_line)
                if match:
                    to_email_address = match.group(1)
                    #print(f"Receivers Email Address is: {to_email_address}")
            
            # Create timeline from after a certain date
            if subjectGlobal == False and emailGlobal == False and dateGlobal == True:
                print('Entered if statement 1\n')
                subject = message.Subject
                body = message.Body
                received_time = message.ReceivedTime
                received_time = message.ReceivedTime
                prompt1 = "Look through this email, make me a timeline of what was agreed to and when. make it concise. for each email, include the date/time, and who sent it. For the content, only include what was either PROMISED or ACCEPTED."
                messagePrompt = subject + "\n" + body + "\n \n" + prompt1

                docString = getGPT3Response(messagePrompt)
                heading = "Subject: " + subject + "\n" + "Received Time: " + str(received_time) + "\n"
                write_to_existing_document(heading, 'timeline.docx')
                write_to_existing_document(docString, 'timeline.docx')
                write_to_existing_document("\n-------------------------New GPT call------------------------\n", 'timeline.docx')
     
                emailsDate = received_time.date().isoformat()
            
                if(emailsDate == dateString):
                    break
                else:
                    print('Date not passed\n')
    return

# Check if co pilot can do it
# Have options to creaete timeline locally or by inputting a document

def is_directory_empty(path):
    return len(os.listdir(path)) == 0

class DialogWindow(QtWidgets.QDialog):
    def __init__(self):
        super().__init__()
        self.ui = Ui_Dialog()
        self.ui.setupUi(self)

        # Connect the clicked signal of checkboxes to the on_checkbox_clicked method
        self.ui.checkBox.clicked.connect(self.on_checkbox_clicked)
        self.ui.checkBox_2.clicked.connect(self.on_checkbox_clicked)

        # Connect the clicked signal of pushButton to the on_pushButton_clicked method
        self.ui.pushButton.clicked.connect(self.on_pushButton_clicked2) #UNCOMMENT THIS LINE TO ENABLE THE BUTTON

    def on_checkbox_clicked(self):
        global emailGlobal, dateGlobal
        if self.ui.checkBox.isChecked():
            dateGlobal = True
        elif self.ui.checkBox_2.isChecked():
            emailGlobal = True
        elif self.ui.checkBox.isChecked() == False:
            dateGlobal = False
        elif self.ui.checkBox_2.isChecked() == False:
            emailGlobal = False
        elif self.ui.checkBox_3.isChecked():
            subjectGlobal = True
        elif self.ui.checkBox_3.isChecked() == False:
            subjectGlobal = False


    def on_pushButton_clicked2(self):
        # Retrieve text from QLineEdit
        text = self.ui.lineEdit.text()
        print("Text from lineEdit:", text)
        if(text != ""):
            emailWG = text
        else:
            emailWG = "default"

        # Retrieve text from QLineEdit
        text2 = self.ui.lineEdit_2.text()
        print("Text from lineEdit_2:", text2)
        if(text2 != ""):
            subjectG = text2
        else:
            subjectG = "default"

        # Retrieve date from QDateEdit
        dateFromG = self.ui.dateEdit.date()
        date_string = dateFromG.toString(QtCore.Qt.DateFormat.ISODate)
        print("Date from dateEdit:", date_string)

        # toString(QtCore.Qt.DateFormat.ISODate)
        #print("Date from dateEdit:", date)

        main(date_string)

########################### End of helper functions ###########################
#
#
#
#
########################### Start of main function ###########################

def main(date_in_String):    
    if os.path.exists('raw_emails.docx'):
        # Delete the file
        os.remove('raw_emails.docx')

    print('timeline.docx was recreated\n')
    if not os.path.exists('timeline.docx'):
        create_document("" ,"timeline.docx")
    elif os.path.exists('timeline.docx'):
        os.remove('timeline.docx')
        create_document("" ,"timeline.docx")

    formTimeline(getEmails(subjectG, emailWG), date_in_String)
    print('\nTimeline has been saved to timeline.docx\n')

    return

########################### End of main function ###########################
#
#
#
#
########################### Start of execution ###########################
if __name__ == "__main__":
    import sys
    app = QtWidgets.QApplication(sys.argv)
    dialog = DialogWindow()
    dialog.show()
    sys.exit(app.exec())
    
print('Starting DisputeLens...\n')
# main()
#print('\nDisputeLens has finished running!')
########################### End of execution ###########################

########################### Code dump ###########################

    # num = 1

    # for message in messages:
    #      # Check if the item is an email message
    #     if hasattr(message, "MessageClass") and message.MessageClass.startswith("IPM.Note"):
    #         if message.ReceivedTime.date() < givenDateFrom:
    #             print("\nEmails after the specified date have been saved to the output folder.")
    #             break

    #         # Get recipients email address
    #         to_email_address = "default"
    #         body2 = message.Body

    #         # Get the line that starts with 'To:'
    #         to_line = next((line for line in body2.split('\n') if line.startswith('To:')), None)

    #         if to_line is not None:
    #             # Use a regular expression to find the email address
    #             match = re.search(r'<(.+)>', to_line)
    #             if match:
    #                 to_email_address = match.group(1)
    #                 #print(f"Receivers Email Address is: {to_email_address}")
    #         # else:
    #         #     print("No 'To:' line found in the email body.")
            
    #         # Select all after a certain date
    #         if subjectGlobal == False and emailGlobal == False and dateGlobal == True:  
    #             # print('Entered if statement 1\n')              
    #             #Get the email subject
    #             subject = message.Subject
    #             #Get the email body
    #             body = message.Body
                
    #             #Get the received date and time
    #             received_time = message.ReceivedTime
    #             #Save the email body and received time to a text file
    #             with open(output_dir / 'all_emails.txt', 'a', encoding='utf-8') as f:
    #                 f.write(f"Email: {num}\n")
    #                 f.write(f"Subject: {subject}\n\n")
    #                 f.write(f"Received Time: {received_time}\n\n")
    #                 f.write(body)
    #                 f.write("\n----------------------\n")
    #             num += 1

    #         # Select all emails with a specific subject
    #         if subjectGlobal == True and emailGlobal == False and dateGlobal == False:
    #             if message.Subject == givenSubject or message.Subject == "Re: " + givenSubject:
    #                 print('Entered if statement 2\n')  
    #                 #Get the email subject
    #                 subject = message.Subject
    #                 #Get the email body
    #                 body = message.Body

    #                 print(body)
    #                 #Get the received date and time
    #                 received_time = message.ReceivedTime
    #                 #Save the email body and received time to a text file
    #                 with open(output_dir / 'all_emails.txt', 'a', encoding='utf-8') as f:
    #                     f.write(f"Email: {num}\n")
    #                     f.write(f"Subject: {subject}\n\n")
    #                     f.write(f"Received Time: {received_time}\n\n")
    #                     f.write(body)
    #                     f.write("\n----------------------\n")
    #                 num += 1
    #                 if message.Subject == "Re: " + givenSubject:
    #                     break

    #         if subjectGlobal == False and emailGlobal == True and dateGlobal == False:
    #             print('Entered if statement 3\n')  
         
    #             if emailWith == to_email_address:                    
    #                 print('Entered the second level function\n\n')
    #                 subject = message.Subject
    #                 #Get the email body
    #                 body = message.Body
    #                 #Get the received date and time
    #                 received_time = message.ReceivedTime
    #                 #Save the email body and received time to a text file
    #                 with open(output_dir / 'all_emails.txt', 'a', encoding='utf-8') as f:
    #                     f.write(f"Email: {num}\n")
    #                     f.write(f"Subject: {subject}\n\n")
    #                     f.write(f"Received Time: {received_time}\n\n")
    #                     f.write(body)
    #                     f.write("\n----------------------\n")
    #                 num += 1

    ########################### End of code dump ###########################